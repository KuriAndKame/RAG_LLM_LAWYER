from docx import Document as DocxDocument
from src.core.models.document import Document


class DocxLoader:
    def load(self, file_path: str) -> tuple[Document, str]:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        document = Document(
            filename=file_path.split('/')[-1],
            source_path=file_path
        )
        return document, text
